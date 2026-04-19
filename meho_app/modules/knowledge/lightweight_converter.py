# SPDX-License-Identifier: AGPL-3.0-only
# Copyright (c) 2026 evoila Group
"""
Lightweight document conversion without PyTorch/Docling.

Uses pymupdf4llm for PDF text/layout, pdfplumber for table extraction,
RapidOCR for scanned PDF OCR, python-docx for DOCX, and BeautifulSoup for HTML.
No GPU required. ~250MB total vs Docling's 2-4GB.

Activated when MEHO_FEATURE_USE_DOCLING=false.
"""

from __future__ import annotations

import contextlib
import io
import tempfile
from dataclasses import dataclass, field
from typing import Any

from meho_app.core.otel import get_logger
from meho_app.modules.knowledge.chunking import TextChunker

logger = get_logger(__name__)


@dataclass
class PageData:
    """Per-page text and metadata."""

    page_number: int  # 1-based
    text: str  # Page text as markdown
    tables: list[str] = field(default_factory=list)  # Markdown-formatted tables


@dataclass
class LightweightDocument:
    """Thin document representation replacing DoclingDocument."""

    markdown: str  # Full document as markdown
    name: str  # Original filename
    page_count: int  # Number of pages (0 for non-PDF)
    pages: list[PageData] = field(default_factory=list)  # Per-page data


SUPPORTED_MIME_TYPES: frozenset[str] = frozenset(
    {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "text/html",
    }
)


class LightweightDocumentConverter:
    """CPU-only document converter matching DoclingDocumentConverter interface.

    Provides the same 3-method interface (convert_file, get_full_text,
    chunk_document) so callers can swap between Docling and lightweight
    paths transparently.
    """

    def __init__(self, max_tokens: int = 512, ocr_enabled: bool = False) -> None:
        self._max_tokens = max_tokens
        self._ocr_enabled = ocr_enabled
        self._chunker = TextChunker(max_tokens=max_tokens)

    def convert_file(self, file_bytes: bytes, filename: str, mime_type: str) -> LightweightDocument:
        """Convert file bytes to a LightweightDocument.

        Args:
            file_bytes: Raw file content.
            filename: Original filename.
            mime_type: MIME type of the file.

        Returns:
            Parsed LightweightDocument.

        Raises:
            ValueError: If the MIME type is unsupported.
        """
        if mime_type == "application/pdf":
            return self._convert_pdf(file_bytes, filename)
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return self._convert_docx(file_bytes, filename)
        elif mime_type == "text/html":
            return self._convert_html(file_bytes, filename)
        else:
            raise ValueError(f"Unsupported MIME type: {mime_type}")

    def get_full_text(self, doc: LightweightDocument) -> str:
        """Export document as markdown text.

        Args:
            doc: Parsed LightweightDocument.

        Returns:
            Full document text in markdown format.
        """
        return doc.markdown

    def chunk_document(
        self, doc: LightweightDocument, chunk_prefix: str = ""
    ) -> list[tuple[str, dict[str, Any]]]:
        """Chunk a LightweightDocument with heading enrichment.

        Delegates to TextChunker.chunk_document_with_structure() for
        markdown heading detection, then adds page_numbers key to match
        DoclingDocumentConverter's output shape.

        Args:
            doc: Parsed LightweightDocument to chunk.
            chunk_prefix: Optional context prefix prepended to each chunk.

        Returns:
            List of (enriched_text, context_dict) tuples where context
            includes heading_stack, page_numbers, and document_name.
        """
        chunks = self._chunker.chunk_document_with_structure(
            doc.markdown, doc.name, detect_headings=True
        )

        results: list[tuple[str, dict[str, Any]]] = []
        for text, context in chunks:
            # Add page_numbers key to match DoclingDocumentConverter output shape
            context["page_numbers"] = []

            # Prepend chunk_prefix if provided
            if chunk_prefix:
                text = chunk_prefix + "\n\n" + text

            results.append((text, context))

        return results

    def _convert_pdf(
        self, file_bytes: bytes, filename: str
    ) -> LightweightDocument:  # NOSONAR (cognitive complexity)
        """Convert PDF using pymupdf4llm text + pdfplumber tables + RapidOCR fallback.

        Pipeline:
        1. pymupdf4llm extracts per-page markdown with heading hierarchy
        2. pdfplumber extracts structured tables per page
        3. Tables are appended to each page's markdown output
        4. RapidOCR handles scanned/image-only pages (if ocr_enabled)
        5. Pages are assembled into a single markdown document

        Args:
            file_bytes: Raw PDF bytes.
            filename: Original filename.

        Returns:
            LightweightDocument with full markdown and per-page data.
        """
        import pymupdf4llm

        # pymupdf4llm needs a file path, so write to temp file
        temp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(file_bytes)
                temp_path = tmp.name

            # Step 1: Extract per-page markdown via pymupdf4llm
            md_pages = pymupdf4llm.to_markdown(
                temp_path,
                show_progress=False,
                page_chunks=True,
            )
            # md_pages is a list of dicts with 'metadata' and 'text' keys
            # metadata has 'page' (0-based page index)

            logger.info(
                "pymupdf4llm_extraction_complete",
                filename=filename,
                pages=len(md_pages),
            )

            # Step 2: Extract tables via pdfplumber
            tables_by_page = self._extract_tables_pdfplumber(file_bytes)

            # Step 3: Merge tables into page text and apply OCR fallback
            page_data_list: list[PageData] = []
            page_texts: list[str] = []

            for i, page_dict in enumerate(md_pages):
                page_text = page_dict.get("text", "")
                page_number = i + 1  # 1-based

                # Step 3a: OCR fallback for near-empty pages
                if self._ocr_enabled and len(page_text.strip()) < 50:
                    ocr_text = self._ocr_page(file_bytes, i)
                    if ocr_text:
                        logger.info(
                            "ocr_fallback_applied",
                            filename=filename,
                            page=page_number,
                            ocr_chars=len(ocr_text),
                        )
                        page_text = ocr_text

                # Step 3b: Append pdfplumber tables to page text
                page_tables = tables_by_page.get(i, [])
                table_md_list: list[str] = []
                for table in page_tables:
                    table_md = _table_to_markdown(table)
                    if table_md:
                        table_md_list.append(table_md)

                if table_md_list:
                    page_text = page_text.rstrip() + "\n\n" + "\n\n".join(table_md_list)

                page_data_list.append(
                    PageData(
                        page_number=page_number,
                        text=page_text,
                        tables=table_md_list,
                    )
                )
                page_texts.append(page_text)

            # Step 4: Assemble full document markdown
            full_markdown = "\n\n---\n\n".join(page_texts)

            logger.info(
                "pdf_conversion_complete",
                filename=filename,
                page_count=len(md_pages),
                total_tables=sum(len(t) for t in tables_by_page.values()),
                markdown_length=len(full_markdown),
            )

            return LightweightDocument(
                markdown=full_markdown,
                name=filename,
                page_count=len(md_pages),
                pages=page_data_list,
            )

        finally:
            if temp_path:
                import os

                with contextlib.suppress(OSError):
                    os.unlink(temp_path)

    def _extract_tables_pdfplumber(
        self, file_bytes: bytes
    ) -> dict[int, list[list[list[str | None]]]]:
        """Extract tables from all PDF pages using pdfplumber.

        Args:
            file_bytes: Raw PDF bytes.

        Returns:
            Dict mapping 0-based page index to list of tables.
            Each table is a list of rows, each row is a list of cell values.
        """
        import pdfplumber

        tables_by_page: dict[int, list[list[list[str | None]]]] = {}

        try:
            pdf = pdfplumber.open(io.BytesIO(file_bytes))
            for page_idx, page in enumerate(pdf.pages):
                try:
                    tables = page.extract_tables()
                    if tables:
                        tables_by_page[page_idx] = tables
                except Exception as e:
                    logger.warning(
                        "pdfplumber_table_extraction_failed",
                        page=page_idx + 1,
                        error=str(e),
                    )
            pdf.close()
        except Exception as e:
            logger.warning(
                "pdfplumber_open_failed",
                error=str(e),
            )

        return tables_by_page

    def _ocr_page(self, file_bytes: bytes, page_index: int) -> str:
        """OCR a single PDF page using RapidOCR via PyMuPDF rendering.

        Args:
            file_bytes: Raw PDF bytes.
            page_index: 0-based page index.

        Returns:
            OCR text or empty string on failure.
        """
        try:
            import pymupdf

            doc = pymupdf.open(stream=file_bytes, filetype="pdf")
            page = doc[page_index]
            pix = page.get_pixmap(dpi=300)
            img_bytes = pix.tobytes("png")
            doc.close()

            from rapidocr_onnxruntime import RapidOCR

            ocr = RapidOCR()
            result, _ = ocr(img_bytes)
            if result:
                return "\n".join([line[1] for line in result])
        except Exception as e:
            logger.warning(
                "ocr_fallback_failed",
                page=page_index + 1,
                error=str(e),
            )

        return ""

    def _convert_docx(
        self, file_bytes: bytes, filename: str
    ) -> LightweightDocument:  # NOSONAR (cognitive complexity)
        """Convert DOCX to markdown using python-docx.

        Walks the document body in element order (paragraphs and tables
        interleaved) to preserve the original structure. Heading styles
        are mapped to markdown heading levels.

        Args:
            file_bytes: Raw DOCX bytes.
            filename: Original filename.

        Returns:
            LightweightDocument with full markdown and no page data (DOCX
            has no page concept without rendering).
        """
        import docx  # python-docx

        doc = docx.Document(io.BytesIO(file_bytes))

        # Build a lookup mapping XML element -> paragraph/table object
        # so we can iterate doc.element.body children in document order.
        para_map: dict[int, Any] = {}
        for para in doc.paragraphs:
            para_map[id(para._element)] = para

        table_map: dict[int, Any] = {}
        for table in doc.tables:
            table_map[id(table._element)] = table

        markdown_lines: list[str] = []
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

        for child in doc.element.body:
            tag = child.tag

            if tag == f"{ns}p":
                para = para_map.get(id(child))  # type: ignore[assignment]
                if para is None:
                    continue
                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else ""

                if style_name == "Title":
                    markdown_lines.append(f"# {text}")
                elif style_name.startswith("Heading"):
                    # Extract heading level from style name (e.g. "Heading 1" -> 1)
                    import re

                    match = re.search(r"\d+", style_name)
                    level = int(match.group()) if match else 1
                    level = min(level, 6)  # Cap at h6
                    markdown_lines.append(f"{'#' * level} {text}")
                elif style_name in ("List Bullet", "List Bullet 2", "List Bullet 3"):
                    markdown_lines.append(f"- {text}")
                elif style_name in ("List Number", "List Number 2", "List Number 3"):
                    markdown_lines.append(f"1. {text}")
                else:
                    markdown_lines.append(text)

            elif tag == f"{ns}tbl":
                table = table_map.get(id(child))  # type: ignore[assignment]
                if table is None:
                    continue
                rows: list[list[str | None]] = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(cells)
                if rows:
                    table_md = _table_to_markdown(rows)
                    if table_md:
                        markdown_lines.append(table_md)

        full_markdown = "\n\n".join(markdown_lines)

        logger.info(
            "docx_conversion_complete",
            filename=filename,
            paragraphs=len(doc.paragraphs),
            tables=len(doc.tables),
            markdown_length=len(full_markdown),
        )

        return LightweightDocument(
            markdown=full_markdown,
            name=filename,
            page_count=0,
            pages=[],
        )

    def _convert_html(
        self, file_bytes: bytes, filename: str
    ) -> LightweightDocument:  # NOSONAR (cognitive complexity)
        """Convert HTML to markdown using BeautifulSoup.

        Strips non-content tags (script, style, nav, footer, header),
        then converts structural HTML elements to markdown equivalents.

        Args:
            file_bytes: Raw HTML bytes.
            filename: Original filename.

        Returns:
            LightweightDocument with full markdown and no page data.
        """
        from bs4 import BeautifulSoup

        # Decode bytes: try UTF-8 first, fall back to latin-1
        try:
            html_text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            html_text = file_bytes.decode("latin-1")

        soup = BeautifulSoup(html_text, "lxml")

        # Remove non-content tags
        for tag_name in ("script", "style", "nav", "footer", "header"):
            for tag in soup.find_all(tag_name):
                tag.decompose()

        # Walk structural elements and convert to markdown
        markdown_lines: list[str] = []
        for element in soup.find_all(
            ["h1", "h2", "h3", "h4", "h5", "h6", "p", "pre", "ul", "ol", "table"]
        ):
            if element.name and element.name.startswith("h") and len(element.name) == 2:
                level = int(element.name[1])
                text = element.get_text(strip=True)
                if text:
                    markdown_lines.append(f"{'#' * level} {text}")
            elif element.name == "p":
                text = element.get_text(strip=True)
                if text:
                    markdown_lines.append(text)
            elif element.name == "pre":
                code_text = element.get_text()
                markdown_lines.append(f"```\n{code_text}\n```")
            elif element.name in ("ul", "ol"):
                for i, li in enumerate(element.find_all("li", recursive=False)):
                    prefix = f"{i + 1}. " if element.name == "ol" else "- "
                    markdown_lines.append(f"{prefix}{li.get_text(strip=True)}")
            elif element.name == "table":
                rows: list[list[str | None]] = []
                for tr in element.find_all("tr"):
                    cells: list[str | None] = [
                        td.get_text(strip=True) for td in tr.find_all(["td", "th"])
                    ]
                    rows.append(cells)
                if rows:
                    table_md = _table_to_markdown(rows)
                    if table_md:
                        markdown_lines.append(table_md)

        full_markdown = "\n\n".join(markdown_lines)

        logger.info(
            "html_conversion_complete",
            filename=filename,
            markdown_length=len(full_markdown),
        )

        return LightweightDocument(
            markdown=full_markdown,
            name=filename,
            page_count=0,
            pages=[],
        )


def _get_pdf_page_count(file_bytes: bytes) -> int:
    """Get PDF page count using pypdfium2 (lightweight, header-only read).

    Duplicated from document_converter.py so the lightweight path
    can import it without pulling in Docling.
    """
    import pypdfium2

    pdf = pypdfium2.PdfDocument(io.BytesIO(file_bytes))
    count = len(pdf)
    pdf.close()
    return count


def build_chunk_prefix(
    connector_type: str | None = None,
    connector_name: str | None = None,
    document_summary: str = "",
) -> str:
    """Build the context prefix prepended to each chunk before embedding (D-05).

    Format: "{connector_type} connector ({connector_name}). {summary}"

    Duplicated from document_converter.py so the lightweight path
    can import it without pulling in Docling.
    """
    parts: list[str] = []
    if connector_type:
        connector_context = f"{connector_type} connector"
        if connector_name:
            connector_context += f" ({connector_name})"
        connector_context += "."
        parts.append(connector_context)
    if document_summary:
        parts.append(document_summary)
    return " ".join(parts)


async def generate_document_summary(
    document_text: str,
    connector_type: str | None = None,
    connector_name: str | None = None,
) -> str:
    """Generate a 1-2 sentence document summary for chunk enrichment (D-05).

    Uses the app's configured classifier model (defaults to Sonnet).
    Returns empty string on failure (ingestion continues without summary).

    Duplicated from document_converter.py so the lightweight path
    can import it without pulling in Docling.
    """
    import asyncio

    from pydantic_ai import Agent

    from meho_app.core.config import get_config

    config = get_config()
    model_name = config.classifier_model  # defaults to anthropic:claude-sonnet-4-6

    agent = Agent(
        model_name,
        system_prompt=(
            "Summarize this document in 1-2 sentences. "
            "Focus on what systems, technologies, or procedures it covers. "
            "If a table of contents is present, use it to identify the major topics. "
            "Return ONLY the summary, no explanation."
        ),
    )

    # First 16K chars (~4 pages) to capture title, TOC, and introduction.
    text_preview = document_text[:16000]

    try:
        result = await asyncio.wait_for(agent.run(text_preview), timeout=15.0)
        return str(result.output).strip()
    except (TimeoutError, Exception) as e:
        logger.warning("document_summary_generation_failed", error=str(e))
        return ""  # Fallback: no summary, still use connector context


def _table_to_markdown(table: list[list[str | None]]) -> str:
    """Convert a pdfplumber table to markdown format.

    Args:
        table: List of rows, each row is a list of cell values (str or None).

    Returns:
        Markdown-formatted table string, or empty string if table is empty.
    """
    if not table or len(table) < 1:
        return ""

    # Clean cells: replace None with empty string, strip whitespace
    cleaned = [[(cell or "").strip() for cell in row] for row in table]

    if not cleaned:
        return ""

    # Header row
    header = cleaned[0]
    col_count = len(header)

    lines: list[str] = []
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * col_count) + " |")

    # Data rows
    for row in cleaned[1:]:
        # Pad or truncate row to match header column count
        padded = row[:col_count] + [""] * max(0, col_count - len(row))
        lines.append("| " + " | ".join(padded) + " |")

    return "\n".join(lines)
